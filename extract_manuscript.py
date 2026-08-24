#!/usr/bin/env python3
"""
Extract the manuscript from Architecture-of-Wonder.docx into manuscript-extracted.txt
in the format build-book.py's parser expects: explicit "CHAPTER N" / "PART X" markers
inferred from Word Heading 1/2/3 styles.

Heading-style heuristic for the new "Built for Wonder" merged DOCX:
  - Heading 1 ALL CAPS  -> chapter title  (auto-numbered)
  - Heading 1 Title Case -> part divider  (mapped to PART ONE..FIVE in order)
  - Heading 1 "THE META REVEAL" -> special, emitted verbatim
  - Heading 1 "CONTROLLING THE ROOM" -> interlude, emitted verbatim (no chapter number)
  - Heading 1 "sw" or empty -> junk, skipped
  - Heading 2 / 3 -> emitted verbatim (parser treats them as content / section headers)

Tables: the document body is walked in order (w:p AND w:tbl), so Word tables
are no longer silently dropped. Each recognised table is emitted as build
marker lines that build-book.py already understands:
  - Cue | Line | Type tables (Ch17 Cold Reading Toolkit) ->
        TOOLKIT_NAV, CR_SUMMARY_TABLE (before the first table), then per table
        TOOLKIT_SECTION: <radar category>
        CRT: <cue> | <type>
        <line>                       (one CRT pair per data row)
        TOOLKIT_SECTION_END
  - ERA x VOWEL grids (Ch23 sitcom titles) ->
        GRID: ERA / VOWEL | A | E | I | O | U
        GRIDROW: <era> | <cell> ...   (one per data row)
  - the T1..T4 evidence-tier table (front matter) is intentionally skipped:
    build-book.py renders it from its own TIER_TABLE_HTML block.
Any other table emits nothing and prints a [warn] line so it cannot spill raw
cells into the manuscript unnoticed.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

DOCX_PATH = Path(__file__).parent / "Built-for-Wonder.docx"
OUT_PATH  = Path(__file__).parent / "manuscript-extracted.txt"

PART_ORDINALS = ["ONE", "TWO", "THREE", "FOUR", "FIVE", "SIX", "SEVEN", "EIGHT"]

# Heading-1 strings that are part dividers, not chapter titles.
# Order matches the book's part order. Match is case-insensitive on text only.
PART_TITLES = [
    "Built for Wonder",
    "Reading the Room",
    "The Methods",
    "Performance Craft",
    "Authority, Influence & The Deep Framework",
]

# Heading-1 strings that are special / not chapters
SPECIAL_VERBATIM = {
    "THE META REVEAL",
    "CONTROLLING THE ROOM",   # unnumbered interlude between CH 37 and 38
}

# Headings whose Word-level text is corrupted but which represent a real chapter.
# Maps the corrupted Heading-1 text -> the title we should emit instead.
# (Verified against running-header text inside the DOCX, e.g.
#  "BUILT FOR WONDERCHAPTER 22 — MAKING BETTER PROPLESS MENTALISM".)
CORRUPTED_HEADING_FIXUPS = {
    "sw": "MAKING BETTER PROPLESS MENTALISM",
}

# Junk headings to skip outright (empty H1s)
JUNK_HEADINGS = {""}

# Chapter titles that should be Heading 1 in the DOCX but are stuck on
# 'normal' style. We promote them to chapter breaks when we encounter them
# in the body text. First occurrence only — subsequent occurrences are
# treated as running-header chrome and dropped.
UNSTYLED_CHAPTER_TITLES = {
    "PROPLESS SYSTEMS THAT ACTUALLY WORK",
}


# The six Cue | Line | Type tables in Chapter 17 map, in document order, onto
# the six radar categories. Names must match the _id_map / _SECTION_ICONS keys
# in build-book.py exactly.
TOOLKIT_CATEGORIES = [
    "Appearance",
    "Movement and Posture",
    "Territory and Personal Space",
    "Social Confidence",
    "Cognitive Processing",
    "Emotional Regulation",
]

# Cells the DOCX uses to mean "no entry" in the vowel grids.
GRID_EMPTY_CELLS = {"", "\u2014", "\u2013", "-"}


def iter_body(doc):
    """Yield ('para', Paragraph) / ('table', Table) for the body's direct
    children in document order. doc.paragraphs is exactly the 'para' subset,
    so paragraph handling is unchanged; tables are the only addition."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "para", Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield "table", Table(child, doc)


def cell_lines(cell):
    """Non-empty paragraph texts of a table cell, stripped."""
    return [t for t in (p.text.strip() for p in cell.paragraphs) if t]


def table_rows(tbl):
    """List of rows; each row is a list of cells; each cell a list of lines."""
    return [[cell_lines(c) for c in r.cells] for r in tbl.rows]


def classify_table(rows):
    """Return one of 'toolkit', 'grid', 'tier', or None (unknown)."""
    if not rows:
        return None
    head = [" ".join(c).strip().lower() for c in rows[0]]
    if head == ["cue", "line", "type"]:
        return "toolkit"
    if len(head) == 6 and head[0] == "" and head[1:] == ["a", "e", "i", "o", "u"]:
        return "grid"
    if head and head[0] == "t1":
        return "tier"
    return None


def emit_toolkit_table(rows, category, out_lines):
    """Emit TOOLKIT_SECTION / CRT marker lines for one Cue | Line | Type table."""
    out_lines.append("")
    out_lines.append(f"TOOLKIT_SECTION: {category}")
    out_lines.append("")
    n = 0
    for r in rows[1:]:
        if len(r) < 3:
            continue
        cue  = " ".join(r[0]).replace("|", "/").strip()
        line = " ".join(r[1]).strip()
        typ  = " ".join(r[2]).replace("|", "/").strip()
        if not cue or not line:
            continue
        out_lines.append(f"CRT: {cue} | {typ}")
        out_lines.append(line)
        out_lines.append("")
        n += 1
    out_lines.append("TOOLKIT_SECTION_END")
    out_lines.append("")
    return n


def emit_grid_table(rows, out_lines):
    """Emit GRID / GRIDROW marker lines for one ERA x VOWEL table."""
    head = ["ERA / VOWEL"] + [" ".join(c).strip() for c in rows[0][1:]]
    out_lines.append("")
    out_lines.append("GRID: " + " | ".join(head))
    for r in rows[1:]:
        cells = []
        for c in r:
            lines = [ln for ln in c if ln not in GRID_EMPTY_CELLS]
            cells.append(" / ".join(lines).replace("|", "/"))
        if not any(cells):
            continue
        out_lines.append("GRIDROW: " + " | ".join(cells))
    out_lines.append("")


def is_all_caps(text: str) -> bool:
    """True if string contains at least one letter and all letters are uppercase."""
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)


def normalize(s: str) -> str:
    return s.strip()


def main():
    if not DOCX_PATH.exists():
        sys.exit(f"DOCX not found: {DOCX_PATH}")

    doc = Document(str(DOCX_PATH))

    out_lines = []
    chapter_counter = 0
    part_index = 0  # which PART_ORDINALS slot we are on
    promoted_unstyled = set()  # titles already promoted to chapter once
    in_toc = False  # we drop the entire TOC; build-book.py renders its own

    table_index = 0          # 1-based, in document order
    toolkit_tables = 0       # how many Cue|Line|Type tables seen so far
    toolkit_rows = 0
    grid_tables = 0
    skipped_tables = []      # (index, kind, shape, first cell) for the report

    for kind, item in iter_body(doc):
        if kind == "table":
            table_index += 1
            if in_toc:
                continue
            rows = table_rows(item)
            shape = f"{len(rows)}x{max((len(r) for r in rows), default=0)}"
            first = " ".join(rows[0][0])[:40] if rows and rows[0] else ""
            ttype = classify_table(rows)
            if ttype == "toolkit":
                if toolkit_tables < len(TOOLKIT_CATEGORIES):
                    category = TOOLKIT_CATEGORIES[toolkit_tables]
                else:
                    category = f"Toolkit {toolkit_tables + 1}"
                    print(f"  [warn] extract: more Cue|Line|Type tables than radar "
                          f"categories; table {table_index} emitted as {category!r}")
                if toolkit_tables == 0:
                    out_lines.append("")
                    out_lines.append("TOOLKIT_NAV")
                    out_lines.append("")
                    out_lines.append("CR_SUMMARY_TABLE")
                    out_lines.append("")
                toolkit_rows += emit_toolkit_table(rows, category, out_lines)
                toolkit_tables += 1
            elif ttype == "grid":
                emit_grid_table(rows, out_lines)
                grid_tables += 1
            elif ttype == "tier":
                # Evidence tiers: build-book.py renders TIER_TABLE_HTML itself.
                skipped_tables.append((table_index, "tier (rendered by build-book)", shape, first))
            else:
                skipped_tables.append((table_index, "UNHANDLED", shape, first))
                print(f"  [warn] extract: unhandled table {table_index} ({shape}, "
                      f"first cell {first!r}) emitted nothing")
            continue

        p = item
        style = p.style.name if p.style else ""
        text  = normalize(p.text)

        # TOC skipping — Heading 2 "Contents" opens it, ends at the next
        # Heading 1/2 (typically "How to Read This Book").
        if in_toc:
            if style in ("Heading 1", "Heading 2"):
                in_toc = False
                # fall through to process this heading normally
            else:
                continue

        if style == "Heading 2" and text.strip().lower() == "contents":
            in_toc = True
            continue

        if style == "Heading 1":
            if text in JUNK_HEADINGS:
                continue

            # Corrupted heading -> rewrite to the real title and treat as chapter.
            if text in CORRUPTED_HEADING_FIXUPS:
                text = CORRUPTED_HEADING_FIXUPS[text]
                chapter_counter += 1
                out_lines.append("")
                out_lines.append(f"CHAPTER {chapter_counter}")
                out_lines.append(text)
                out_lines.append("")
                continue

            if text in SPECIAL_VERBATIM:
                out_lines.append("")
                out_lines.append(text)
                out_lines.append("")
                continue

            # Part divider?  Match by exact (case-sensitive) title or by title-case heuristic.
            is_part = False
            matched_part_title = None
            for pt in PART_TITLES:
                if text == pt:
                    is_part = True
                    matched_part_title = pt
                    break

            if is_part:
                if part_index >= len(PART_ORDINALS):
                    # Defensive: more parts than expected, just skip the marker
                    out_lines.append(matched_part_title)
                    continue
                ordinal = PART_ORDINALS[part_index]
                part_index += 1
                out_lines.append("")
                out_lines.append(f"PART {ordinal}")
                out_lines.append(matched_part_title)
                out_lines.append("")
                continue

            # Otherwise treat as a chapter title (whether ALL CAPS or not — H1
            # at this point is a chapter heading)
            chapter_counter += 1
            out_lines.append("")
            out_lines.append(f"CHAPTER {chapter_counter}")
            out_lines.append(text)
            out_lines.append("")
            continue

        if style == "Heading 2":
            # Treat as front-matter section header / content.  Some H2s in this
            # DOCX are TOC-ish ("Contents"), some are real ("How to Read This
            # Book", "ACKNOWLEDGMENTS"). The build-book.py parser looks for
            # exact strings — emit as-is and uppercased for the structural ones.
            upper = text.upper()
            if upper in {"HOW TO READ THIS BOOK", "ACKNOWLEDGMENTS"}:
                out_lines.append("")
                out_lines.append(upper)
                out_lines.append("")
            else:
                # Misc H2 (like emoji-prefixed category headers in CH 17,
                # "ONE WORD TITLES", "Contents"): pass through verbatim.
                out_lines.append(text)
            continue

        if style == "Heading 3":
            # Section sub-headers — pass through verbatim, build script can
            # detect them by Title Case matching later.
            out_lines.append(text)
            continue

        # Normal / Body / unstyled paragraphs.
        # Promote certain unstyled titles to real chapter breaks (first occurrence only).
        if text in UNSTYLED_CHAPTER_TITLES and text not in promoted_unstyled:
            promoted_unstyled.add(text)
            chapter_counter += 1
            out_lines.append("")
            out_lines.append(f"CHAPTER {chapter_counter}")
            out_lines.append(text)
            out_lines.append("")
            continue

        out_lines.append(text)

    # Collapse 3+ consecutive blank lines into 2.
    cleaned = []
    blank_run = 0
    for line in out_lines:
        if line == "":
            blank_run += 1
            if blank_run <= 2:
                cleaned.append(line)
        else:
            blank_run = 0
            cleaned.append(line)

    OUT_PATH.write_text("\n".join(cleaned) + "\n", encoding="utf-8")

    print(f"Wrote {OUT_PATH.name}")
    print(f"  Chapters detected: {chapter_counter}")
    print(f"  Parts emitted:     {part_index}")
    print(f"  Total lines:       {len(cleaned):,}")
    print(f"  Tables walked:     {table_index} "
          f"(toolkit {toolkit_tables} = {toolkit_rows} CRT rows, "
          f"grid {grid_tables}, skipped {len(skipped_tables)})")
    for idx, why, shape, first in skipped_tables:
        print(f"    table {idx}: {why}, {shape}, first cell {first!r}")


if __name__ == "__main__":
    main()
