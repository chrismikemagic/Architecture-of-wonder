#!/usr/bin/env python
"""One-shot: insert the Imbalanced section into Chapter 24 of Built-for-Wonder.docx,
after the Reflex material and before the chapter-closing '· · ·' separator.
Source: 'Imbalanced' by Atlas Brookings (The Intrepid Rogue's Manual of Deception),
an impromptu propless reworking of Marc Paul's 'True Lies'. Reprinted/adapted with
permission from Brookings; Marc Paul credited in prose. Voice rewritten to match chapter.

CLEARANCE: pending Marc Paul approval for the underlying 'True Lies' principle.
To remove cleanly if declined, see CLEARANCES.md (revert the insertion commit)."""
from docx import Document
from docx.shared import Emu

DOCX = "Built-for-Wonder.docx"
HEADER_SIZE = Emu(171450)  # 13.5pt — matches existing Ch24 subsection headers

BLOCKS = [
    ("IMBALANCED", "title"),
    ("“Imbalanced” by Atlas Brookings", "byline"),
    ("Reflex puts one person in front of you and reads the body for a single hidden bit at a time. Imbalanced does the opposite. It hands the deduction to a whole line of people at once, hides the work inside a story, and lets you name a liar, a truth teller, and a murderer with nothing on you but your attention. It is the group-scale version of everything this chapter is built on: you already know more than they think, and the body fills in the rest.", "body"),
    ("The routine is Atlas Brookings’, and he built it on the spot one night when a prop failed him mid-performance and he had to rebuild Marc Paul’s “True Lies” with empty hands. Brookings is emphatic that the heart of it is Marc Paul’s, and so am I. If you perform this, learn Marc’s work and credit him by name.", "body"),
    ("The Effect", "subhead"),
    ("Six people stand in a line and give you a thumbs up. You tell them you are about to stage a murder mystery, and each of them will play a part. Three of them will be liars — not casually dishonest, but pathological, lying about everything. Three will be truth tellers, obsessively honest. The three who choose to be liars quietly turn their thumbs down. The three who stay honest leave their thumbs up. Then, among themselves, the group silently elects a murderer by pointing.", "body"),
    ("You never see any of it. You face away while they sort themselves out. When they are ready, you turn back, pace the line like Poirot, ask a few questions, and announce who lied, who told the truth, and who did it.", "body"),
    ("The Imbalance", "subhead"),
    ("Here is the whole engine, and it is the reason for the name.", "body"),
    ("You go down the line and ask each person a single question: are you the murderer? You do not need a thing beyond their answers.", "body"),
    ("Three people are honest and three are lying, so the answers never split evenly. Suppose the murderer is one of the liars. The three honest people all say no, because they are innocent and honest. The guilty liar also says no, because he lies about being the murderer. That is four noes. The two remaining liars are innocent, so they lie and claim the crime — two yeses. Four against two.", "body"),
    ("Run it the other way. Suppose the murderer is honest. He admits it: one yes. The two other honest people deny it truthfully: two noes. All three liars are innocent and lie about it, claiming the murder: three yeses. Four yeses against two noes.", "body"),
    ("So the split is never even, and the direction of the imbalance hands you two facts at once. An imbalance toward yes means your murderer is honest. An imbalance toward no means your murderer is a liar. And the smaller group — the two who answered against the crowd — is always innocent. Always. Set them aside.", "body"),
    ("To keep the count clean, do not try to track six individual answers in your head. Ask the question and have anyone who claims to be the murderer step forward. Now the imbalance is something you can see at a glance: four standing, two stepped out, or the reverse. You read the direction, you know your murderer’s role, and you know your group of four contains the person you are hunting.", "body"),
    ("Narrowing the Four", "subhead"),
    ("You now know one thing that matters: whether the murderer is a liar or a truth teller. Your job is to find the one person in the group of four whose honesty matches that role. You do it by asking questions you already know the answer to, and watching whether each suspect answers them straight.", "body"),
    ("This is where the chapter’s earlier work pays off. You have been reading these people since they raised their hands. You noted who is right-handed and who is not from the hand they gave you the thumbs up with, the side they part their hair, the wrist that carries the watch, the way a belt is looped. You noted who is wearing a wedding ring. None of it looked like method. All of it is ammunition now.", "body"),
    ("The handedness question. Pick two of your four suspects you know for a fact are right-handed, and work the theme: “It is obvious from the damage that our killer is right-handed. You — are you right-handed?” A truth teller says yes. A liar says no. If the two you ask answer the same way, they share a role and you dismiss them, leaving two suspects. If they answer opposite to each other, those two are your final pair and you can dismiss everyone else. Either way you have learned who lies and who does not, measured against a fact you already held.", "body"),
    ("A different question, so the pattern hides. Once you have asked about hands, do not ask about hands again. Reach for another fact you already own. If you know one of your remaining suspects is married, build the frame and use it: “The killer is a monster — I pity anyone close to him. You there. Are you single?” The married man who says he is single has just lied to you, and if your murderer is a liar, you are done.", "body"),
    ("The age trap. Ask a suspect how old they are, then immediately ask what year they were born. A truth teller answers both without friction. A liar who gave you a false age now has to do quiet arithmetic to keep the year consistent, and the pause — or the math that does not add up — is the tell. Wrap it in the story: “It took strength to do this. I think we are looking for someone young. How old are you?” Then, on the heels of it: “So what year were you born?” The hesitation is your answer.", "body"),
    ("The Foolproof Out", "subhead"),
    ("There is a foolproof out built into the structure, and you should know it even if you rarely reach for it. Because of how the imbalance works, you already know the roles of the two people in the smaller group — the innocents. If you ever lose the thread, point at two members of the larger group and ask one of your known innocents whether one of them is the killer. If you asked a liar and they say yes, the two you pointed at are both innocent. If they say no, one of the pair you indicated is your target.", "body"),
    ("It even sets up a laugh. An innocent liar will have claimed the murder himself during the first round. When he then fingers someone else, you can hold his eye and ask, “Helped you do it, did they?” Send him back to his seat, innocent but troubled.", "body"),
    ("The Last Suspect Comes Free", "subhead"),
    ("You will usually never question your final suspect at all. By the time you have read the others, the last person’s role and guilt are fixed by elimination — you know their answer before they give it. Leave them unquestioned and the logic puzzle disappears, or, if you want the bigger finish, ask them to answer only in their mind, then tell them what they are thinking. The deduction becomes a mind-read.", "body"),
    ("Scaling It Down", "subhead"),
    ("Imbalanced works with any even number of people, but more bodies mean more rounds of narrowing questions. After enough performances Brookings settled on four participants rather than six. With four you ask one less round and reach the reveal faster, and for most rooms that is the stronger version. Start with six while you are learning the shape of it, then tighten to four once the imbalance reads itself to you on sight.", "body"),
    ("Why It Belongs Here", "subhead"),
    ("Everything in Reflex is one body, one bit at a time. Imbalanced is the same skill turned outward: you arrive already knowing more than the room believes you could, and the only thing you add in the moment is observation under a story good enough that no one feels watched. The murder frame is doing exactly what the which-hand frame does — giving the audience something to enjoy while you read what their bodies were never trying to hide.", "body"),
]


def main():
    d = Document(DOCX)
    ps = d.paragraphs

    # Anchor: the Ch24-closing separator '· · ·' that precedes
    # 'You can get a lot of information when you use a which hand...'.
    anchor = None
    for i, p in enumerate(ps):
        if p.text.strip() in ("· · ·", "• • •") and i + 1 < len(ps) \
                and ps[i + 1].text.strip().lower().startswith("you can get a lot of information"):
            anchor = p
            break
    if anchor is None:
        raise SystemExit("ERROR: could not find the Ch24-closing separator anchor.")

    ref = anchor._p
    for text, kind in BLOCKS:
        new_p = d.add_paragraph()
        new_p.style = ps[0].style  # 'normal'
        run = new_p.add_run(text)
        if kind == "subhead":
            run.bold = True
            run.font.size = HEADER_SIZE
        ref.addprevious(new_p._p)

    d.save(DOCX)
    print(f"Inserted {len(BLOCKS)} paragraphs before the Ch24-closing separator.")


if __name__ == "__main__":
    main()
