#!/usr/bin/env python
"""Apply the verified mechanical typo/spacing fixes from the 2026-07-05 editorial
audit to Built-for-Wonder.docx.

Source: scratchpad/confirmed-fixes.json (83 fixes that passed the auditor's
safe_to_auto_fix flag AND an adversarial verifier). Six were EXCLUDED after a
char-level review because they were rewrites / would introduce errors / are
header-separator style choices:
  #61 truncated rewrite starting with '...'
  #77 drops leading 'T' (drop-cap misread) -> introduces error
  #79 injects a newline (extraction artifact, two joined headers)
  #78, #81 introduce em/en dashes into section headers (author's call)
  #82 removes a possibly-intentional repeated 'Walk in certain.'

Application is run-aware: if the target string lives inside one run, only that
run is edited (all inline formatting preserved). If it spans runs, the edit is
applied only when every run in the paragraph shares the same bold/italic/
underline state (collapse is then lossless); otherwise the fix is SKIPPED and
reported so nothing with mixed formatting is clobbered. Any fix whose exact
text is not found, or is found in more than one paragraph, is also SKIPPED and
reported. Backs up the DOCX first.
"""
import json
import shutil
import sys
from datetime import datetime
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent
DOCX = ROOT / "Built-for-Wonder.docx"
BACKUPS = ROOT / "backups"
FIXES_JSON = Path(r"C:\Users\Chris\AppData\Local\Temp\claude\C--Users-Chris\0c56a7b7-b5ef-4868-a4b8-fd1a776bdf0b\scratchpad\confirmed-fixes.json")

EXCLUDE = {61, 77, 78, 79, 81, 82}


def all_paragraphs(doc):
    """Body paragraphs + paragraphs inside every table cell."""
    ps = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                ps.extend(cell.paragraphs)
    return ps


def run_fmt(r):
    return (bool(r.bold), bool(r.italic), bool(r.underline))


def try_apply(p, old, new):
    """Return ('applied'|'skip-mixedfmt'|'nomatch', detail)."""
    runs = p.runs
    if not runs:
        return ('nomatch', 'no runs')
    joined = "".join(r.text or "" for r in runs)
    idx = joined.find(old)
    if idx < 0:
        return ('nomatch', 'string not in paragraph')
    if joined.find(old, idx + 1) >= 0:
        return ('nomatch', 'string appears >1x in paragraph')
    end = idx + len(old)
    # locate run span
    pos = 0
    start_run = end_run = None
    local_start = local_end = None
    for ri, r in enumerate(runs):
        rt = r.text or ""
        rstart, rend = pos, pos + len(rt)
        if start_run is None and rstart <= idx < rend:
            start_run = ri
            local_start = idx - rstart
        if rstart < end <= rend:
            end_run = ri
            local_end = end - rstart
        pos = rend
    if start_run is None or end_run is None:
        return ('nomatch', 'could not map run span')
    if start_run == end_run:
        r = runs[start_run]
        r.text = r.text[:local_start] + new + r.text[local_end:]
        return ('applied', 'single-run')
    # spans runs: only safe if all involved runs share formatting
    fmts = {run_fmt(runs[ri]) for ri in range(start_run, end_run + 1)}
    if len(fmts) > 1:
        return ('skip-mixedfmt', f'spans runs {start_run}-{end_run} with mixed formatting')
    # lossless collapse across the involved runs: put replacement text in the
    # first involved run, keep the prefix/suffix outside the [old] span.
    prefix = (runs[start_run].text or "")[:local_start]
    suffix = (runs[end_run].text or "")[local_end:]
    runs[start_run].text = prefix + new + suffix
    for ri in range(start_run + 1, end_run + 1):
        runs[ri].text = ""
    return ('applied', f'collapsed runs {start_run}-{end_run} (uniform fmt)')


def main():
    fixes = json.load(open(FIXES_JSON, encoding="utf-8"))
    todo = [(i, f) for i, f in enumerate(fixes) if i not in EXCLUDE]
    print(f"Loaded {len(fixes)} verified fixes; applying {len(todo)} (excluded {sorted(EXCLUDE)})")

    BACKUPS.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    shutil.copy2(DOCX, BACKUPS / f"Built-for-Wonder-pre-audit-fixes-{ts}.docx")
    shutil.copy2(DOCX, BACKUPS / "Built-for-Wonder-pre-audit-fixes.docx")
    print(f"Backed up DOCX -> backups/Built-for-Wonder-pre-audit-fixes.docx")

    d = Document(str(DOCX))
    paras = all_paragraphs(d)

    applied, skipped = [], []
    for i, f in todo:
        old = f.get("exact_quote", "")
        new = f.get("proposed_fix", "")
        ch = f.get("chapter", "?")
        if not old or new == "" or old == new:
            skipped.append((i, ch, "empty/identical", old))
            continue
        # find candidate paragraphs
        cands = [p for p in paras if old in ("".join(r.text or "" for r in p.runs))]
        if len(cands) == 0:
            skipped.append((i, ch, "not found", old))
            continue
        if len(cands) > 1:
            skipped.append((i, ch, f"found in {len(cands)} paragraphs", old))
            continue
        status, detail = try_apply(cands[0], old, new)
        if status == 'applied':
            applied.append((i, ch, detail, old, new))
        else:
            skipped.append((i, ch, detail, old))

    d.save(str(DOCX))

    print(f"\n=== APPLIED {len(applied)}/{len(todo)} ===")
    for i, ch, detail, old, new in applied:
        print(f"  #{i} [{ch}] {old[:48]!r} -> {new[:48]!r}")
    print(f"\n=== SKIPPED {len(skipped)} ===")
    for i, ch, reason, old in skipped:
        print(f"  #{i} [{ch}] ({reason}) {old[:60]!r}")

    # persist an apply-log for the report / snapshot
    log = {
        "applied": [{"idx": i, "chapter": ch, "detail": d_, "was": o, "now": n} for i, ch, d_, o, n in applied],
        "skipped": [{"idx": i, "chapter": ch, "reason": r, "text": o} for i, ch, r, o in skipped],
        "excluded_from_autoapply": sorted(EXCLUDE),
    }
    logpath = Path(r"C:\Users\Chris\AppData\Local\Temp\claude\C--Users-Chris\0c56a7b7-b5ef-4868-a4b8-fd1a776bdf0b\scratchpad\apply-log.json")
    json.dump(log, open(logpath, "w", encoding="utf-8"), indent=2, ensure_ascii=False)
    print(f"\nApply-log -> {logpath}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
