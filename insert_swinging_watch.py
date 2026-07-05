#!/usr/bin/env python
"""One-shot: insert 'Where the Swinging Watch Comes From' into Chapter 19
(HOW HYPNOSIS REALLY WORKS) of Built-for-Wonder.docx.

Placed right after the 'The Alpha Shift' section (anchor = its closing
paragraph) and before the 'Down to the Cell Level' header, grouping it with the
chapter's other 'how the receptive state is entered' material. The section
explains that rhythmic light pulsing across closed eyes (sun through trees,
streetlights sweeping past a back-seat passenger) echoes the eye activity of
REM/hypnagogic drift, and that this is the real mechanism behind the swinging
watch. It pays off the chapter's own 'no swinging watch' setup and its dream
section.

This is CHRIS'S OWN material (his idea), not sourced from Atlas Brookings, so it
ships in BOTH editions and gets NO brookings_manifest entry.

Voice: no em dashes; uneven contrasts and standalone imperatives only (no
mirror-image contrast pairs, no deny-rename-command triads); REM glossed inline.
"""
import shutil
from datetime import datetime
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent
DOCX = ROOT / "Built-for-Wonder.docx"
BACKUPS = ROOT / "backups"

ANCHOR = "Together they produce a measurable shift in the receiving architecture"
NEXT_HEADER = "Down to the Cell Level"

BLOCKS = [
    ("Where the Swinging Watch Comes From", "header"),
    ("Think about the last time you rode in the back of a car with your eyes closed. Not asleep, just resting, letting the trip happen to you. If it was daytime, the sun was strobing through the trees and the buildings, sweeping across your closed lids in a slow, uneven rhythm. If it was night, the streetlights slid past one after another, doing the same thing. And somewhere in the middle of all that gentle flicker you got heavy, warm, pleasantly and almost helplessly sleepy.", "body"),
    ("There is a real mechanism under that heaviness, and it is the same one a hypnotist reaches for.", "body"),
    ("When light pulses across the eyes at a slow, repeating rhythm, even through closed lids, it stimulates the visual system in a way that echoes what the brain does at the edge of sleep. Rapid eye movement, the stage of sleep where most dreaming happens, is named for exactly that, the eyes flicking in rhythmic bursts while the outside world is shut out. Rhythmic light and rhythmic eye activity are part of the signature the brain reads as drifting off. Feed that signature back in from the outside, gently and over and over, and you nudge the nervous system toward the same drowsy, inward, hypnagogic state, the half-dreaming edge you cross on the way into sleep. The vigilant part of attention has nothing new to track. Every pulse is the same as the last one, so it stops standing guard and starts to settle.", "body"),
    ("You already met this earlier in the chapter, coming from the other direction. The dream is the brain building a whole world from the inside with the eyes closed. The flicker is a doorway to that world, propped slightly open by nothing more exotic than moving light.", "body"),
    ("Now look at the swinging watch again.", "body"),
    ("For a hundred years the pocket watch on a chain has been the cartoon shorthand for hypnosis, and most performers roll their eyes at it. But it was never nonsense. A small bright object, swung slowly back and forth, hands the eyes a single monotonous thing to follow. The metal catches the light and throws a soft glint on every pass. The eyes track it side to side, close to the way they move in sleep, and the light pulses with the swing. Hold it a little above eye level and the eyes tire faster and want to close on their own. Everything the watch does on purpose, a car window does by accident.", "body"),
    ("The watch was never magic in the metal. It was rhythm and light doing to a nervous system exactly what rhythm and light always do. Once you see that, you can put the watch away and keep the principle. A steady voice with a repeating cadence, a slow gesture repeated until it stops being watched, a single point to rest the eyes on, every one of them reaches for the same door.", "body"),
]


def main():
    if not DOCX.exists():
        raise SystemExit(f"DOCX not found: {DOCX}")

    BACKUPS.mkdir(exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d-%H%M%S")
    for name in ("Built-for-Wonder-pre-swinging-watch.docx",
                 f"Built-for-Wonder-pre-swinging-watch-{ts}.docx"):
        shutil.copy2(DOCX, BACKUPS / name)
    print(f"Backed up DOCX -> backups/Built-for-Wonder-pre-swinging-watch.docx (+timestamped)")

    d = Document(str(DOCX))
    ps = d.paragraphs

    target_idxs = [i for i, p in enumerate(ps) if ANCHOR in p.text]
    if len(target_idxs) != 1:
        raise SystemExit(f"ERROR: expected exactly 1 anchor paragraph, found {len(target_idxs)}")
    ti = target_idxs[0]
    target = ps[ti]

    # Sanity: the immediate next non-empty paragraph should be the NEXT_HEADER.
    # (Confirms we are inserting at the Alpha Shift / Down-to-the-Cell-Level seam.)
    following = next((p.text.strip() for p in ps[ti + 1:] if p.text.strip()), "")
    if following != NEXT_HEADER:
        raise SystemExit(f"ERROR: expected next header {NEXT_HEADER!r}, found {following!r}")

    normal_style = ps[0].style
    new_ps = []
    for text, _kind in BLOCKS:
        np = d.add_paragraph()
        np.style = normal_style
        np.add_run(text)
        new_ps.append(np._p)

    ref = target._p
    for elem in reversed(new_ps):
        ref.addnext(elem)

    d.save(str(DOCX))
    print(f"Inserted 'Where the Swinging Watch Comes From' ({len(BLOCKS)-1} body paras) after The Alpha Shift.")


if __name__ == "__main__":
    main()
