#!/usr/bin/env python
"""One-shot: insert 'The Pre-Show That Looks Like a Failure' as the final section of
Chapter 27 (Pre-Show), after the 'Lines That Redirect...' material and before the
chapter-closing '· · ·'.

Source: the 'Fail.' pre-show approach from 'Train Tracking' by Atlas Brookings,
reprinted/adapted with permission. Method-agnostic: Train Tracking's proprietary
letter 'Tracks' are NOT reproduced. Voice rewritten to match the chapter.
Brookings material — excluded from the no-Brookings build."""
from docx import Document
from docx.shared import Emu

DOCX = "Built-for-Wonder.docx"
HEADER_SIZE = Emu(171450)

ANCHOR = "universal behavior, not a specific catch"

BLOCKS = [
    ("THE PRE-SHOW THAT LOOKS LIKE A FAILURE", "title"),
    ("Everything in this chapter so far hides the pre-show by making it forgettable — an innocent, logistical moment the participant files under nothing. Atlas Brookings taught me the opposite move, and it is one of the cleverest pre-show ideas I know. Instead of making the earlier moment forgettable, you make it a failure. You let them remember it perfectly, because what they remember is you trying and not getting it.", "body"),
    ("Here is the shape of it. Before the show, you take a participant aside and run the exact read you intend to perform on stage later — the same opening words, the same instructions, the same procedure. You do the whole thing. By the time you are finished, you already know everything you need. But you reveal none of it. You shrug, you look faintly disappointed in yourself, and you say, “I’m sorry — I couldn’t read you at all. Tell you what, let me come back to this later in the show.”", "body"),
    ("That is the entire method. You fail on purpose.", "body"),
    ("Think about what that failure buys you. The standard worry with pre-show is that the participant, or someone sitting near them, will later trace the effect back to that quiet conversation before the doors opened. Here there is nothing to trace. They did not watch you succeed. They watched you try and come up empty. Nobody guards the memory of a thing that did not work. When you return to them on stage and get it, there is no earlier success for them to connect it to — only an earlier failure, which makes the later hit feel like something you finally broke through to, not something you arranged.", "body"),
    ("It also disarms the participant completely. They are not surprised when you circle back, because you told them you might. They are not suspicious, because in their own experience you genuinely did not get it the first time. And they are quietly invested — they want to see whether you can do now what you couldn’t before. You have turned the one person most able to expose you into the one person rooting hardest for the miracle.", "body"),
    ("A few practical notes from Brookings that keep it reliable.", "body"),
    ("Set up more than one. If you run the failed read on three people before the show, you give yourself margin. If one of them is half-checked-out and somehow drifts to a new thought, you simply leave them in their seat and work the other two. Brookings says this has never actually gone wrong for him; the insurance just costs you nothing.", "body"),
    ("You barely have to carry anything in your head. Depending on the method you use, you usually only need to remember each person’s original seed — the single word or thought everything else grew out of. From that one anchor, the rest of what they pictured reconstructs itself when you need it.", "body"),
    ("Cue the callback in plain sight. When you come back to them on stage, the line that does the work is some version of, “What’s interesting is what you were thinking about before the show even started.” That sentence quietly tells the participant which moment you are reaching back to, so they lock onto the earlier experience — while the audience hears nothing but a mind reader widening his claim. If someone ever looks lost, the word tonight pulls them back: “You’re sure you weren’t seeing a lion tonight?” reads as eerie to the room and as a gentle reminder to them.", "body"),
    ("The reason this belongs beside everything else in this chapter is that it solves the same problem from the far side. The rest of pre-show works by making the secret moment vanish. This works by making it visible and harmless. Both leave the participant able to tell the exact truth about what happened and still have it sound impossible. “He talked to me before the show and couldn’t read me at all — and then later he told me everything I had been thinking.” Let them say that. It is the best testimony you could ask for.", "body"),
]


def main():
    d = Document(DOCX)
    ps = d.paragraphs
    hits = [p for p in ps if ANCHOR in p.text]
    if len(hits) != 1:
        raise SystemExit(f"ERROR: expected exactly 1 anchor, found {len(hits)}")
    ref = hits[0]._p

    new_ps = []
    for text, kind in BLOCKS:
        np = d.add_paragraph()
        np.style = ps[0].style  # 'normal'
        run = np.add_run(text)
        if kind == "title":
            run.bold = True
            run.font.size = HEADER_SIZE
        new_ps.append(np._p)

    # Insert all new paragraphs immediately after the anchor, preserving order.
    for elem in reversed(new_ps):
        ref.addnext(elem)

    d.save(DOCX)
    print(f"Inserted {len(BLOCKS)} paragraphs after the Ch27 anchor.")


if __name__ == "__main__":
    main()
