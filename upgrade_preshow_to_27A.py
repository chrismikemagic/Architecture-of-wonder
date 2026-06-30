#!/usr/bin/env python
"""Upgrade the 'The Pre-Show That Looks Like a Failure' section into a standalone
Chapter 27A interlude beat, mirroring the 7A/21A/37A pattern.

Adds (after the existing title):
  - a hook epigraph (renders as a bold-italic quote, like 21A's hook)
  - the A-chapter chrome: tier/category badges + 'CHAPTER 27A — ...' marker
    (the build SKIPS these lines; they exist for source/Word identity only)
And adds (after the last body paragraph, before the chapter's '· · ·'):
  - a closing key-read aphorism (renders as a paragraph, like 21A's closer)

The manifest anchor ('THE PRE-SHOW THAT LOOKS LIKE A FAILURE' -> '· · ·') stays
valid, so the no-Brookings edition still strips the whole beat cleanly."""
from docx import Document

DOCX = "Built-for-Wonder.docx"

TITLE = "THE PRE-SHOW THAT LOOKS LIKE A FAILURE"
LAST_BODY_MATCH = "the best testimony you could ask for"

HOOK = "“The safest secret is the one they watched you fail to find.”"
# A-chapter chrome (all skipped by the build; present for source identity).
CHROME = [
    "T1T2",
    "SIGNAL CONFIDENCE TIERS",
    "AM",
    "OBSERVATION CATEGORIES",
    "CHAPTER 27A — THE PRE-SHOW THAT LOOKS LIKE A FAILURE",
]
KEY_READ = "Fail first, and the miracle has no past for them to search."


def _new_para(d, ref_p, text, after=True):
    np = d.add_paragraph()
    np.style = d.paragraphs[0].style
    np.add_run(text)
    if after:
        ref_p.addnext(np._p)
    else:
        ref_p.addprevious(np._p)
    return np._p


def main():
    d = Document(DOCX)
    ps = d.paragraphs

    title_p = next((p for p in ps if p.text.strip() == TITLE), None)
    if title_p is None:
        raise SystemExit("ERROR: title paragraph not found.")
    last_body_p = next((p for p in ps if LAST_BODY_MATCH in p.text), None)
    if last_body_p is None:
        raise SystemExit("ERROR: last body paragraph not found.")

    # After the title: hook, then chrome (in order). Insert reversed via addnext.
    after_title = [HOOK] + CHROME
    for text in reversed(after_title):
        _new_para(d, title_p._p, text, after=True)

    # After the last body paragraph: the key-read closing line.
    _new_para(d, last_body_p._p, KEY_READ, after=True)

    d.save(DOCX)
    print("Upgraded section to Chapter 27A beat: +hook, +chrome, +key-read.")


if __name__ == "__main__":
    main()
