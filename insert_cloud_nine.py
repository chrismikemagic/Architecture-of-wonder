#!/usr/bin/env python
"""One-shot: insert the Cloud Nine section into Chapter 25 of Built-for-Wonder.docx,
after 'The Red Dwarf' and before the chapter's closing '· · ·' separator.
Source: 'Cloud Nine' by Atlas Brookings (The Intrepid Rogue's Manual of Deception),
reprinted/adapted with permission. Year engine seeded with Andrew Brown; start/middle/
end-of-month refinement by Michael Murray. Voice rewritten to match the chapter."""
from docx import Document
from docx.shared import Emu

DOCX = "Built-for-Wonder.docx"
HEADER_SIZE = Emu(171450)  # 13.5pt — matches existing Ch25 subsection headers

# Block list: (text, kind)  kind in {title, byline, subhead, body}
BLOCKS = [
    ("CLOUD NINE", "title"),
    ("“Cloud Nine” by Atlas Brookings", "byline"),
    ("Everything in this chapter so far has been built to reach one thing: the sign. But the sign is only a slice of a larger number. A person’s zodiac is decided entirely by the date they were born, which means if you can read the whole birthdate, you already have the sign sitting inside it. This final method does exactly that, with no props, no anagrams, and no moment where the room hears you miss. Every calculation happens in the spectator’s head. You are only ever comparing their numbers against one number of your own.", "body"),
    ("Atlas Brookings calls it Cloud Nine, and he gave me permission to teach it here. It began as an email exchange between Atlas and Andrew Brown about the hidden mathematics of a birth year, and the start, middle, and end-of-month idea you will use to pin the day belongs to Michael Murray. I have rewritten the handling in my own words, but the engine is theirs. Credit them when you perform it.", "body"),
    ("The whole system rests on a quirk of the century we live in. Almost everyone you will ever read was born in a year that begins 19 or 20. That single fact is doing all of the work.", "body"),
    ("The Frame", "subhead"),
    ("You are going to present this as numerology rather than astrology. Numerology gives you a reason to ask for arithmetic, and arithmetic is what hides the method.", "body"),
    ("“For a long time numerologists treated your star sign as a rough sketch. They believed the numbers in your birth date told a sharper story. I want to try the version of that I actually trust. None of this needs to be said out loud. Just do the math in your head and answer with comparisons, not numbers.”", "body"),
    ("That instruction — comparisons, not numbers — is the engine. The spectator never tells you a value. They only ever tell you bigger, smaller, or by how much.", "body"),
    ("Reading the Year", "subhead"),
    ("Have them picture their birth year as four digits: two on the inside, two on the outside. For someone born in 1987, the inner pair is 9 and 8, the outer pair is 1 and 7.", "body"),
    ("“Multiply the two inner digits together. If you land on a two-digit number, add those digits and keep going until you have a single digit.”", "body"),
    ("Here is the secret hiding in plain sight. Any pair of inner digits from a 19-something year includes that 9, and any number built on a nine reduces back to nine. So for anyone born last century, the inner number is always 9. You know it before they finish the math. For someone born this century, the inner pair includes a 0, so their inner number is always 0 — more on the young ones in a moment.", "body"),
    ("“Now multiply the two outer digits together and reduce the same way.”", "body"),
    ("The outer pair is the first digit of the year times the last digit. For a 19-something year the first digit is 1, and anything times 1 is itself. So their outer number is simply the last digit of their birth year, handed to you.", "body"),
    ("You do not ask for it. You ask them to compare.", "body"),
    ("“Don’t tell me either number. Just tell me — is the first number you made larger or smaller than the second? And by how much?”", "body"),
    ("You already know the inner number is 9. If they say it was larger by two, then the outer number is 9 minus 2, which is 7. The last digit of their birth year is 7. You now hold 19_7, and the only thing missing is the decade.", "body"),
    ("You get the decade by looking at them. A glance tells you whether 1987 or 1977 or 1997 fits, the same way you have been estimating throughout this whole chapter. Call the third digit and the year is complete. Treat the nine as their dominant number out loud, and tell them it is the only number they need to keep in mind. Every other number will come and go.", "body"),
    ("The Day-and-Month Number", "subhead"),
    ("“Take the number of your birth month — January is one, December is twelve — and add it to the day you were born. Reduce that to a single digit.”", "body"),
    ("April the 18th gives 4 plus 18, which is 22, which reduces to 4. Again, you do not want the value, you want the comparison.", "body"),
    ("“Is that number the same as your dominant number, higher, or lower? And by how much?”", "body"),
    ("It cannot be higher than nine, but you play as though it could. If they say lower by five, then their day-and-month number is 9 minus 5, which is 4. Hold onto that four. You will hunt for it shortly.", "body"),
    ("Somewhere in here, casually and well before the reveal, ask whether they were born toward the start, the middle, or the end of the month. It draws no attention as a stray question and it saves you later.", "body"),
    ("Reading the Month", "subhead"),
    ("“One more. Take your birth-month number again and add it to your dominant number. Reduce it to a single digit. This one is your happiness factor.”", "body"),
    ("The happiness-factor line is cover. What you are actually doing is adding nine to the month. And because of how nine behaves, that sum reduces straight back to the month itself for any month from four through nine. Ask for the comparison — lower by five means 9 minus 5, which is 4 — and you read the month directly. April.", "body"),
    ("The only soft spot is the low end. Months one, two, and three share their reduced value with October, November, and December. When you land on a one, two, or three, you pump to break the tie: ask whether they have had their birthday yet this year, or talk about the weather, the holidays, or the length of their month. One question settles it.", "body"),
    ("Pinning the Exact Day", "subhead"),
    ("You know the month value and the day-and-month number. To find the first candidate date, run this: 9 minus the month value, plus the day-and-month number.", "body"),
    ("For April, that is 9 minus 4 plus 4. Work left to right: 9 minus 4 is 5, plus 4 is 9. The 9th of April gives you that day-and-month value. Add nine and add nine again, and you have the only other candidates: the 18th and the 27th. If your starting date comes out higher than nine, subtract nine as well as add it.", "body"),
    ("Now the question you already asked pays off. They told you middle of the month. The 9th is early, the 27th is late, the 18th sits in the middle. If two candidates ever fall in the same third of the month, one will be odd and one will be even — a quiet “this isn’t an even date, is it?” separates them.", "body"),
    ("April 18th, 1987. You never heard a number, and you never guessed in front of anyone.", "body"),
    ("The Young Ones", "subhead"),
    ("For a spectator born this century, the inner pair of their year contains a zero, so their dominant number is 0 instead of 9, and you compare everything against zero rather than nine. The outer pair is 2 times the last digit, so reduce it as they go: even results halve, odd results start at five and climb by one for each step past one. You will rarely need it, but it is there.", "body"),
    ("Why It Belongs Here", "subhead"),
    ("This is the cleanest answer to the worry that runs under every bold method in this chapter — what happens when I am wrong? With Cloud Nine, there is no public moment to be wrong in. The spectator does private arithmetic and reports comparisons. If a read of their age or month is off, you adjust inside the same conversation, because nothing has been committed to the room. The tells that confirm you are on the right line are the ones from Part Two; if you need to refresh them, the green, yellow, and red signals on Chris Michael’s tell table apply here exactly as they did to the Repeat It Ploy earlier in this chapter.", "body"),
    ("The sign was always the small version of this. Here you take the whole date, and the sign comes free.", "body"),
]


def main():
    d = Document(DOCX)
    ps = d.paragraphs

    # Locate anchor: the chapter-closing separator '· · ·' that precedes
    # 'The zodiac is a frame...'. Insert our blocks right before it.
    anchor = None
    for i, p in enumerate(ps):
        if p.text.strip() in ("· · ·", "• • •") and i + 1 < len(ps) \
                and ps[i + 1].text.strip().lower().startswith("the zodiac is a frame"):
            anchor = p
            break
    if anchor is None:
        raise SystemExit("ERROR: could not find the chapter-closing separator anchor.")

    ref = anchor._p
    for text, kind in BLOCKS:
        new_p = d.add_paragraph()              # appended at end for now
        new_p.style = ps[0].style              # 'normal'
        run = new_p.add_run(text)
        if kind == "subhead":
            run.bold = True
            run.font.size = HEADER_SIZE
        # title and byline keep Red Dwarf's plain formatting (matches existing)
        ref.addprevious(new_p._p)              # move into place before separator

    d.save(DOCX)
    print(f"Inserted {len(BLOCKS)} paragraphs before the chapter-closing separator.")


if __name__ == "__main__":
    main()
