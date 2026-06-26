#!/usr/bin/env python
"""One-shot: insert 'The Three Eras of Method' into Chapter 1 of Built-for-Wonder.docx.
The Ch1 closer is a single paragraph holding two sentences split by a line break:
  '...some VERY good methods too.\\n\\nNow we can formally begin... Let's get started.'
We split that paragraph and drop the Three Eras section between the halves, so the
flow is: thesis wrap -> historical lineage -> 'Now we can formally begin'.

Source: 'The Three Eras of Mentalism' essay from 'Decoherence' by Atlas Brookings,
reprinted/adapted with permission. Voice rewritten into Chris's; framed to reinforce
the chapter thesis. Brookings material — excluded from the no-Brookings build."""
from docx import Document

DOCX = "Built-for-Wonder.docx"

PART_B = "Now we can formally begin where I originally intended. Let’s get started."

BLOCKS = [
    ("THE THREE ERAS OF METHOD", "title"),
    ("There is a way to look at the entire history of this craft that lands on the same point Colin made, and it comes from Atlas Brookings. He divides mentalism into three eras, and what they really chart is method slowly getting out of the way.", "body"),
    ("The first era was mechanical. From the early nineteen-hundreds, performers like Annemann and Dunninger got hold of a thought by physically getting hold of it. You wrote something down, a billet was switched, a peek was stolen, a nail writer did its quiet work. The secret lived in the hands. The method was a piece of machinery hidden under the table, and the whole art was running it without being caught.", "body"),
    ("Then came the control era. By the middle of the century, people like Hummer, Hoy, and Goldstein walked away from sleight and built structures instead — stacks, anagrams, the Gilbreath principle, a little binary math. The performer stopped reaching into your pocket and started fishing in your mind, narrowing and nodding, taking confirmation you never noticed you were giving. The machinery moved out of the hands and into the procedure.", "body"),
    ("What we are in now Brookings calls the conceptual era. No billet, no peek, no fishing. The participant is handed real freedom — think of anything, change your mind, write nothing down — and the performer reads what is left behind: a hesitation, a flicker of doubt, the shape a free choice leaves as it is made. Train Tracking is one of these. So is most of what you are about to learn in this book.", "body"),
    ("Notice the direction of travel. Each era handed the participant more freedom and asked the method to do less. We went from taking the thought, to steering the thought, to simply being in the room while someone has one. A hundred years of innovation, spent making the method quieter.", "body"),
    ("Follow that line to its end and you arrive exactly where this chapter began. The further the method recedes, the more clearly you can see what was carrying the moment the whole time. It was never the billet or the stack or the clever bit of structure. It was a person paying close attention to another person, in a way that felt like something real was happening. That is what Colin watched on that stage, and it is what the best of this craft has always quietly been. The history of our method is the story of it learning to step out of the way of that.", "body"),
]


def main():
    d = Document(DOCX)
    ps = d.paragraphs

    # Find the Ch1 closing paragraph that contains both sentences.
    target = None
    for p in ps:
        if "formally begin where I originally intended" in p.text and \
           "ultimately building toward" in p.text:
            target = p
            break
    if target is None:
        raise SystemExit("ERROR: could not find the Ch1 closing paragraph.")

    full = target.text
    idx = full.find("Now we can formally begin")
    part_a = full[:idx].strip()

    # Rewrite the target paragraph to hold only part A.
    for r in list(target.runs):
        r._element.getparent().remove(r._element)
    target.add_run(part_a)

    # Build the new paragraphs (Three Eras blocks + part B) and insert them
    # immediately after the (now part-A-only) target paragraph, in order.
    ref = target._p
    new_ps = []
    for text, kind in BLOCKS:
        np = d.add_paragraph()
        np.style = ps[0].style  # 'normal'
        np.add_run(text)
        new_ps.append(np._p)
    # part B as its own paragraph, after the Three Eras section
    pb = d.add_paragraph()
    pb.style = ps[0].style
    pb.add_run(PART_B)
    new_ps.append(pb._p)

    # Insert all new paragraphs right after `target`, preserving order.
    for elem in reversed(new_ps):
        ref.addnext(elem)

    d.save(DOCX)
    print(f"Split Ch1 closer and inserted {len(BLOCKS)} Three Eras paragraphs + closing line.")


if __name__ == "__main__":
    main()
